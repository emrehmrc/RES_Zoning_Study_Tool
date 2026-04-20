"""
Generate a comprehensive User Manual for the Renewable Energy Zoning Dashboard.
Output: User_Manual.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Global styles ───────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Heading styles
for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)  # dark blue

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 3'].font.size = Pt(13)
doc.styles['Heading 4'].font.size = Pt(11)

# Helper: add a styled table
def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    # Rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def add_note(doc, text, bold_prefix="Note: "):
    """Add a highlighted note paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(bold_prefix)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.italic = True


def add_tip(doc, text):
    """Add a tip paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run("Tip: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x6B, 0x3F)
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.italic = True


def add_placeholder(doc, caption="[Screenshot placeholder]"):
    """Add a placeholder box for manual screenshot insertion."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"  {caption}  ")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for _ in range(6):
    p.add_run("\n")

title_run = p.add_run("Renewable Energy Zoning Dashboard")
title_run.bold = True
title_run.font.size = Pt(28)
title_run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = p2.add_run("User Manual")
sub_run.font.size = Pt(20)
sub_run.font.color.rgb = RGBColor(0x3D, 0x7E, 0xAA)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.add_run("\n\n")
ver = p3.add_run("Version 3.0")
ver.font.size = Pt(14)
ver.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = p4.add_run("April 2026")
date_run.font.size = Pt(12)
date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

add_placeholder(doc, "[Company logo placeholder]")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS PLACEHOLDER
# ════════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
p = doc.add_paragraph()
p.add_run("(Right-click \u2192 Update Field to generate the Table of Contents after inserting screenshots.)")
p.runs[0].font.italic = True
p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# Insert a TOC field
p_toc = doc.add_paragraph()
run_toc = p_toc.add_run()
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run_toc._r.append(fldChar1)
run_toc2 = p_toc.add_run()
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
run_toc2._r.append(instrText)
run_toc3 = p_toc.add_run()
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
run_toc3._r.append(fldChar2)
run_toc4 = p_toc.add_run("(Table of Contents \u2014 press Ctrl+A then F9 to update)")
run_toc4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run_toc5 = p_toc.add_run()
fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run_toc5._r.append(fldChar3)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  SECTION 1 — INTRODUCTION
# ════════════════════════════════════════════════════════════════════
doc.add_heading("1. Introduction", level=1)

doc.add_heading("1.1 Purpose and Scope", level=2)
doc.add_paragraph(
    "The Renewable Energy Zoning Dashboard is a GIS-based web application designed "
    "for identifying, evaluating, and ranking optimal sites for renewable energy projects. "
    "The platform supports three distinct project modes:"
)
bullets = [
    ("Solar PV Zoning", "Identifies and scores land parcels suitable for photovoltaic solar installations based on solar irradiation, terrain, land use, and infrastructure proximity."),
    ("On-Shore Wind Zoning", "Evaluates inland areas for wind turbine deployment using wind resource data, terrain constraints, environmental exclusions, and grid connectivity."),
    ("Off-Shore Wind Zoning", "Analyzes maritime Exclusive Economic Zones (EEZ) for offshore wind farms, incorporating bathymetry, seabed substrate, marine constraints, port proximity, and subsea infrastructure."),
]
for title, desc in bullets:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + " \u2014 ")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    "The application guides the user through a structured four-step analytical pipeline \u2014 "
    "from defining the study area and generating a spatial grid, through multi-layer raster "
    "and vector analysis, weighted scoring with hard exclusion constraints, spatial clustering, "
    "and finally financial viability assessment. Each step builds upon the results of the "
    "previous one, creating a comprehensive site suitability analysis."
)

doc.add_heading("1.2 Technology Stack", level=2)
doc.add_paragraph(
    "The dashboard is built on a modern full-stack architecture orchestrated with Docker Compose:"
)
add_table(doc,
    ["Component", "Technology", "Description"],
    [
        ["Frontend", "Next.js 14, React 18, Tailwind CSS, Leaflet", "Interactive UI with map visualization at port 3000"],
        ["Backend", "Django 4.2, Django REST Framework, Gunicorn", "REST API serving analysis endpoints at port 8000"],
        ["GIS Engines", "GeoPandas, Rasterio, GDAL 3.11, NetworkX, Shapely", "Spatial analysis, raster/vector processing, graph-based clustering"],
        ["Deployment", "Docker Compose", "Two-container setup with shared data volumes"],
    ],
    col_widths=[3, 5.5, 7.5]
)

doc.add_heading("1.3 Four-Step Analytical Pipeline", level=2)
doc.add_paragraph(
    "The entire workflow is organized into four sequential tabs in the dashboard interface, "
    "each corresponding to one pipeline step:"
)

add_table(doc,
    ["Step", "Dashboard Tab", "Purpose", "Key Output"],
    [
        ["1", "\U0001f4d0 1. Gridization", "Define the study area boundary and divide it into uniform rectangular grid cells", "Grid of cells covering the selected region"],
        ["2", "\U0001f3af 2. Layer Calculation", "Analyze each grid cell against multiple GIS raster and vector layers (distances, coverage, statistics)", "Per-cell metrics for each configured layer"],
        ["3", "\U0001f4c8 3. Scoring", "Assign weighted scores to each cell; apply hard exclusion constraints; kV layer weights feed into Step 4 connection scoring", "FINAL_GRID_SCORE per cell (0\u2013100)"],
        ["4", "\U0001f9e9 4. Cluster & Aggregation", "Group adjacent high-scoring cells into project clusters; score transmission connections; calculate CAPEX, LCOE, and payback period", "Ranked clusters with financial feasibility metrics"],
    ],
    col_widths=[1, 3.5, 5, 5.5]
)

doc.add_paragraph(
    "Results are visualized on an interactive Leaflet map at each stage, and can be downloaded "
    "as CSV files for further analysis in external tools such as QGIS or Excel."
)

doc.add_heading("1.4 Session Management", level=2)
doc.add_paragraph(
    "When a project mode is selected, the backend assigns a unique Session ID (UUID) stored in "
    "the browser's localStorage as dashboard_session_id. "
    "This identifier is attached to every subsequent API request via the X-Session-ID HTTP header. "
    "All intermediate results \u2014 grid data, analysis outputs, scoring tables \u2014 are persisted "
    "server-side as serialized files (pickle format) under a session-specific directory "
    "(temp/sessions/<uuid>/). Session metadata is stored as session_meta.json alongside "
    "the pickle files."
)
add_note(doc, "Sessions are temporary. If the backend container is restarted or the session is reset, all intermediate results will be lost. Always download important results before closing the application.")

doc.add_heading("1.5 Asynchronous Operations", level=2)
doc.add_paragraph(
    "The three computation-intensive steps \u2014 raster/vector analysis (Step 2), scoring (Step 3), "
    "and cluster analysis (Step 4) \u2014 run as asynchronous background tasks. When triggered, the "
    "backend immediately returns a task_id. The frontend continuously polls "
    "GET /api/task/<task_id>/progress/ to retrieve the current progress (0\u2013100%), status message, "
    "and completion state. A ProcessingOverlay component displays an animated spinner and the "
    "latest progress message during execution. This design allows large analyses to run without "
    "blocking the browser."
)

doc.add_heading("1.6 Getting Started", level=2)
doc.add_paragraph(
    "To launch the dashboard, ensure Docker is installed and run the following command "
    "from the project root directory:"
)
p = doc.add_paragraph()
run = p.add_run("    docker compose up")
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph("Once the containers are running:")
steps = [
    "Open a web browser and navigate to http://localhost:3000",
    "The Landing Page will appear with three project mode cards",
    "Click on a project mode to begin (e.g., \"Select Solar PV Zoning\")",
    "You will be redirected to the main Dashboard view with four analysis tabs",
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  SECTION 2 — SOLAR PV ZONING STUDY
# ════════════════════════════════════════════════════════════════════
doc.add_heading("2. Solar PV Zoning Study \u2014 Detailed Software Functionalities", level=1)

doc.add_paragraph(
    "This section provides a comprehensive walkthrough of the Solar PV Zoning mode, "
    "covering every screen, button, input field, and interactive element from the moment "
    "the application is opened. While this guide focuses on the Solar PV workflow, "
    "the On-Shore and Off-Shore Wind modes follow the same four-tab structure with "
    "project-specific layers and parameters noted where applicable."
)

# ── 2.0 Landing Page ───────────────────────────────────────────────
doc.add_heading("2.0 Landing Page \u2014 Project Mode Selection", level=2)

add_placeholder(doc, "[Screenshot: Landing page with three project cards]")

doc.add_paragraph(
    "Upon navigating to the application URL, the user is presented with the Landing Page. "
    "This page serves as the entry point and project mode selector."
)

doc.add_heading("2.0.1 Page Layout", level=3)
doc.add_paragraph(
    "The Landing Page features a dark gradient background (slate tones) with the following elements:"
)
elements = [
    ("Main Title", "\"Renewable Energy Zoning Dashboard\" \u2014 displayed prominently at the top center."),
    ("Subtitle", "\"Select a project mode to begin analysis\" \u2014 guides the user to choose a mode."),
    ("About Button", "Located in the top-left corner. Labeled \"About\" \u2014 opens an informational modal describing the platform's purpose, pipeline steps, and technical architecture."),
    ("OST Logo", "Displayed in the top-right corner of the page."),
    ("Project Mode Cards", "Three large selection cards arranged horizontally (responsive layout)."),
]
for title, desc in elements:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + ": ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("2.0.2 Project Mode Cards", level=3)
doc.add_paragraph("Each card contains:")
items = [
    "A project-specific icon image (Solar panel, wind turbine on land, or offshore turbine)",
    "The project mode title (e.g., \"Solar PV Zoning\")",
    "A brief list of key analysis capabilities",
    "A selection button (e.g., \"Select Solar PV Zoning\")",
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

add_table(doc,
    ["Card", "Title", "Key Capabilities", "Button Color"],
    [
        ["Solar PV", "Solar PV Zoning", "Solar PV Potential Analysis; Slope & Terrain & Constraints; Proximity to Transmission Lines", "Orange"],
        ["On-Shore Wind", "On-Shore Wind Zoning", "Wind Resource & Potential; Turbine Specific Suitability; Environmental & Social Constraints", "Dark Blue"],
        ["Off-Shore Wind", "Off-Shore Wind Zoning", "Wind Resource & Potential; Turbine Specific Suitability; Marine Constraints & Seabed Analysis", "Cyan/Blue"],
    ],
    col_widths=[2.5, 3.5, 6, 2.5]
)

doc.add_paragraph(
    "Clicking a card initiates a backend call (POST /api/project/select/) to create a new session "
    "with the selected project type. A loading spinner appears on the card during this process. "
    "Upon success, the user is automatically redirected to the main Dashboard."
)

doc.add_heading("2.0.3 About Modal", level=3)
doc.add_paragraph(
    "The About modal provides detailed information organized into sections: "
    "Overview, How It Works (4-Step Pipeline), Solar PV Mode, On-Shore Wind Mode, "
    "Off-Shore Wind Mode, Financial Analysis, Technical Architecture, and Data & Outputs. "
    "The modal can be closed by clicking the X button or clicking outside the modal area."
)

doc.add_page_break()

# ── 2.0.4 Main Dashboard Layout ────────────────────────────────────
doc.add_heading("2.0.4 Main Dashboard Layout", level=2)

add_placeholder(doc, "[Screenshot: Dashboard overview with sidebar and tab bar visible]")

doc.add_paragraph(
    "After selecting Solar PV mode, the Dashboard page loads. The interface consists of:"
)

doc.add_heading("Header Bar", level=3)
items = [
    ("Color Accent Bar", "A thin colored line at the very top of the page (orange for Solar PV, dark blue for On-Shore Wind, cyan for Off-Shore Wind) provides instant visual identification of the active project mode."),
    ("Project Icon", "The project-specific icon is displayed on the left side of the header."),
    ("Application Title", "Shown next to the icon, pulled from the project configuration."),
    ("Mode Badge", "Displays the active project type (e.g., \"Mode: Solar\") in the header's right area."),
    ("Switch Mode Button", "Labeled \"Switch Mode\" \u2014 prompts for confirmation, resets the current session entirely, clears the localStorage session ID, and returns to the Landing Page for a new project selection."),
]
for title, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + ": ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("Sidebar (Left Panel)", level=3)
doc.add_paragraph(
    "A fixed-width sidebar (256 pixels) on the left side displays the project status at a glance:"
)
items = [
    ("Status Title", "Shows the project type followed by \"Status\" (e.g., \"Solar Status\")."),
    ("Step Indicators", "Four status rows \u2014 Grid, Layers, Scoring, Clusters \u2014 each showing a checkmark (\u2705 completed) or hourglass (\u23f3 pending) icon. When a step is completed, it also displays the count of items (e.g., \"1,250 cells\" for Grid, \"5 layer(s)\" for Layers, cluster count for Clusters)."),
    ("Reset Project Button", "A red button labeled \"Reset Project\" at the bottom of the sidebar. Prompts for confirmation before clearing all project data while keeping the same project mode active (calls POST /api/project/reset/ with keep_project_type: true)."),
]
for title, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + ": ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("Tab Bar (Main Content Area)", level=3)
doc.add_paragraph(
    "The main content area occupies the remaining space to the right of the sidebar. "
    "At the top, a horizontal tab bar provides navigation between the four pipeline steps:"
)
add_table(doc,
    ["Tab #", "Label", "Pipeline Step"],
    [
        ["1", "\U0001f4d0 1. Gridization", "Define study area & create grid"],
        ["2", "\U0001f3af 2. Layer Calculation", "Raster and vector analysis across all layers"],
        ["3", "\U0001f4c8 3. Scoring", "Weighted scoring, exclusion constraints, kV connection weights"],
        ["4", "\U0001f9e9 4. Cluster & Aggregation", "Clustering, connection scoring, financial analysis"],
    ],
    col_widths=[1.5, 4, 9]
)
doc.add_paragraph(
    "The active tab is highlighted with a distinct background color and bottom border accent. "
    "Inactive tabs are shown in muted gray tones and become highlighted on hover."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  2.1 TAB 1 — GRIDIZATION
# ════════════════════════════════════════════════════════════════════
doc.add_heading("2.1 Tab 1 \u2014 Gridization", level=2)

add_placeholder(doc, "[Screenshot: Gridization tab \u2014 Generate New Grid mode]")

doc.add_paragraph(
    "The Gridization tab is the first step in the analytical pipeline. Its purpose is to define "
    "the geographic study area and subdivide it into a uniform grid of rectangular cells. "
    "Each cell will later be individually analyzed, scored, and potentially clustered."
)

doc.add_heading("2.1.1 Mode Selection", level=3)
doc.add_paragraph(
    "At the top of the tab, two toggle buttons allow the user to choose between:"
)
items = [
    ("\U0001f30d Generate New Grid", "Create a fresh grid by selecting a geographic boundary and specifying cell dimensions. This is the primary workflow."),
    ("\U0001f4e4 Upload Existing Grid", "Import a previously created grid as a CSV file. Useful for re-running analysis on a previously defined study area."),
]
for title, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + ": ")
    run.bold = True
    p.add_run(desc)

# Generate New Grid
doc.add_heading("2.1.2 Generate New Grid", level=3)
doc.add_paragraph(
    "When \"Generate New Grid\" is selected, the interface splits into two columns:"
)

doc.add_heading("Left Column: Boundary Definition", level=4)
doc.add_paragraph(
    "This section determines the geographic area that will be divided into grid cells."
)

doc.add_paragraph("For Solar PV and On-Shore Wind modes:")
items = [
    ("Select Country", "Choose from a dropdown of European countries derived from the NUTS (Nomenclature of Territorial Units for Statistics) dataset at the national level (LEVL_CODE = 0). The dropdown is populated automatically from the GET /api/countries/ endpoint."),
    ("Albania Sub-National Drill-Down", "When Albania is selected from the country dropdown, two additional dropdowns appear: (1) Region \u2014 populated from GET /api/albania/regions/ listing all ADM1 administrative regions; (2) District \u2014 populated from GET /api/albania/districts/?region=... filtered by the selected region, listing ADM2 districts. Selecting a district uses the precise district boundary rather than the entire country boundary for grid generation."),
]
for title, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + ": ")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    "For Off-Shore Wind mode, the boundary source is fixed to EEZ (Exclusive Economic Zone) selection. "
    "A dropdown is populated with available European maritime zones from GET /api/eez-zones/."
)

add_note(doc, "Once a boundary is selected, an interactive map preview appears below the form, showing the boundary polygon and a real-time grid overlay.")

doc.add_heading("Right Column: Grid Parameters", level=4)
doc.add_paragraph("The grid dimensions depend on the project type:")

doc.add_heading("Solar PV:", level=4)
add_table(doc,
    ["Parameter", "Input Type", "Range", "Step", "Default", "Description"],
    [
        ["Grid Width (m)", "Number", "100 \u2013 10,000 m", "100 m", "1,000 m", "Horizontal dimension of each grid cell in meters"],
        ["Grid Height (m)", "Number", "100 \u2013 10,000 m", "100 m", "1,000 m", "Vertical dimension of each grid cell in meters"],
    ],
    col_widths=[3, 2, 2.5, 1.5, 1.5, 4]
)

doc.add_heading("On-Shore Wind:", level=4)
add_table(doc,
    ["Parameter", "Input Type", "Range", "Step", "Default", "Description"],
    [
        ["Turbine Diameter (m)", "Number", "20 \u2013 200 m", "10 m", "200 m", "Rotor diameter; grid cell Width = 3 \u00d7 D, Height = 5 \u00d7 D"],
    ],
    col_widths=[3, 2, 2.5, 1.5, 1.5, 4]
)

doc.add_heading("Off-Shore Wind:", level=4)
add_table(doc,
    ["Parameter", "Input Type", "Range", "Step", "Default", "Description"],
    [
        ["Turbine Diameter (m)", "Number", "20 \u2013 500 m", "10 m", "200 m", "Rotor diameter; grid cell Width = 3 \u00d7 D, Height = 5 \u00d7 D"],
    ],
    col_widths=[3, 2, 2.5, 1.5, 1.5, 4]
)

doc.add_paragraph(
    "For wind projects, the grid cell size is calculated automatically based on standard turbine spacing rules: "
    "the width is set to 3 times the turbine diameter, and the height to 5 times the diameter. "
    "For example, a 200m diameter turbine yields a 600m \u00d7 1000m grid cell. "
    "An information box below the input field displays the computed dimensions in real time."
)
add_note(doc, "Input validation is applied in real time. If a value falls outside the allowed range, the input border turns red and an error message is displayed. The value is automatically clamped to the valid range when the field loses focus.")

doc.add_heading("2.1.3 Map Preview", level=3)

add_placeholder(doc, "[Screenshot: Map preview showing boundary polygon and grid overlay]")

doc.add_paragraph(
    "When a boundary (country, EEZ zone, or Albania district) is selected, an interactive Leaflet "
    "map appears below the form. The map provides:"
)
items = [
    ("Boundary Polygon", "The selected region boundary displayed as a blue polygon with semi-transparent fill."),
    ("Grid Overlay", "A visualization of the grid cells as colored lines, computed in real time. If the projected cell count exceeds 300,000 cells, the grid overlay is suppressed with a warning message \u2014 the grid will still be created successfully but cannot be rendered at that density."),
    ("Base Map Toggle", "Two radio buttons in the top-right corner allow switching between \"Street\" (OpenStreetMap) and \"Satellite\" (ArcGIS World Imagery) base maps."),
    ("Layer Toggles", "Checkboxes to show/hide the boundary and grid overlay independently."),
    ("Grid Information", "A small info panel in the bottom-left corner displays the grid cell size and the estimated total number of cells."),
]
for title, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + ": ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("2.1.4 Create Grid Button", level=3)
doc.add_paragraph(
    "The \"\U0001f680 Create Grid\" button (blue, full-width) at the bottom of the parameters section "
    "initiates the grid generation process. When clicked:"
)
steps = [
    "A processing overlay appears with an animated spinner and the message \"Creating grid cells...\"",
    "The backend FastGridEngine reprojects the boundary to Web Mercator (EPSG:3857), calculates the bounding box, and generates rectangular cells",
    "Only cells whose centroids fall within the boundary polygon are retained",
    "Each cell is assigned a unique cell_id and its geometry is stored as WKT (Well-Known Text)",
    "The overlay disappears and results are displayed on the map and in the results panel",
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_heading("2.1.5 Upload Existing Grid", level=3)
doc.add_paragraph(
    "When the \"Upload Existing Grid\" mode is selected, a file drop area appears. "
    "The user can drag and drop a CSV file onto it, or click to open the file browser. "
    "The file is automatically uploaded upon selection via POST /api/grid/upload/. "
    "The CSV must contain at least the following columns:"
)
items = ["cell_id \u2014 a unique identifier for each grid cell", "wkt \u2014 the Well-Known Text geometry representation of each cell (EPSG:3857)"]
for item in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item.split(" \u2014 ")[0])
    run.bold = True
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    p.add_run(" \u2014 " + item.split(" \u2014 ")[1])

doc.add_heading("2.1.6 Grid Results", level=3)

add_placeholder(doc, "[Screenshot: Grid results with success message and preview table]")

doc.add_paragraph("After successful grid creation, the following elements appear:")
items = [
    ("Success Message", "A green banner displays the confirmation message (e.g., \"Grid created successfully\") along with the total number of cells generated and the grid cell dimensions."),
    ("Download CSV Button", "Labeled \"\U0001f4e5 Download CSV\" \u2014 calls GET /api/grid/download/ and exports the full grid dataset as a comma-separated file (grid.csv, UTF-8 encoding)."),
    ("Preview Table", "A scrollable AnalysisResultsTable showing the generated grid with pagination, sortable columns, and per-column text filters. Clicking a row focuses that cell on the map. Displays columns including cell_id, bounding coordinates (left, top, right, bottom), WKT geometry, and center point coordinates."),
    ("Map with Grid Cells", "The Leaflet map updates to display the actual generated grid cells aligned to real spatial coordinates."),
]
for title, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + ": ")
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  2.2 TAB 2 — LAYER CALCULATION
# ════════════════════════════════════════════════════════════════════
doc.add_heading("2.2 Tab 2 \u2014 Layer Calculation", level=2)

add_placeholder(doc, "[Screenshot: Layer Calculation tab with layers configured]")

doc.add_paragraph(
    "The Layer Calculation tab is the second pipeline step. "
    "Here, the user configures which GIS raster (and optionally vector) layers to analyze, "
    "and how each layer should be processed for every grid cell. "
    "This step transforms raw spatial data into per-cell numerical metrics."
)

doc.add_heading("2.2.1 Prerequisites", level=3)
doc.add_paragraph(
    "This tab requires a grid to have been created in Step 1. If no grid data is available, "
    "a warning message is displayed: \"No grid data available. Complete Step 1 first.\""
)

doc.add_heading("2.2.2 Add New Layer", level=3)
doc.add_paragraph(
    "The \"Add New Layer\" section provides a form to add spatial layers for analysis. "
    "Two modes are available via toggle buttons:"
)

doc.add_heading("Predefined List Mode", level=4)
doc.add_paragraph(
    "This mode presents a dropdown of predefined layer names specific to the active project type. "
    "When a predefined layer is selected, the system automatically assigns the appropriate analysis "
    "modes based on the project configuration. Layers already added are removed from the dropdown."
)
doc.add_paragraph("For Solar PV, the predefined layers are organized into categories:")

add_table(doc,
    ["Category", "Layer Names", "Default Analysis Modes"],
    [
        ["Infrastructure \u2013 Transmission Lines", "Distance to 110kV Line, Distance to 220kV Line, Distance to 400kV Line", "distance"],
        ["Infrastructure \u2013 Substations", "Distance to 110kV Substation, Distance to 220kV Substation, Distance to 400kV Substation", "distance"],
        ["Land Use & Environment", "Agricultural Areas, Forest, Urban/Residential/Industrial, Military, Protected Habitats", "distance, coverage"],
        ["Natural Resources", "Energy Sources, Hydrography, Mineral Resources", "distance, coverage"],
        ["Risk & Climate", "Natural Risk Zones, Slope (%), Solar Irradiation, Temperature", "Varies: coverage, or min/max/mean"],
        ["Transportation", "Transport Networks", "distance"],
    ],
    col_widths=[3, 5.5, 3.5]
)

add_note(doc, "On-Shore Wind and Off-Shore Wind modes have their own predefined layer lists. On-Shore adds Wind Speed (max/mean/min), Altitude (max), and Airports. Off-Shore uses Wind Speed, Bathymetry, Slope, Sea bed (shapefile), Ports, Subsea Cables, and other marine layers.")

doc.add_heading("Custom Layer Mode", level=4)
doc.add_paragraph(
    "In Custom mode, the user can define a new layer from scratch:"
)
items = [
    ("Custom Layer Name", "A free-text input for the layer name (e.g., \"My Custom Layer\")."),
    ("Analysis Modes", "A set of toggle pill buttons for selecting one or more modes. Available modes are: distance, coverage, mean, max, min, median, std. Multiple modes can be selected simultaneously. Selected modes appear in blue."),
    ("Target Pixel Value", "Shown only when \"distance\" or \"coverage\" is among the selected modes. This numeric input (0\u2013255) specifies which raster pixel value to target for distance calculation or coverage percentage computation. Default: 1."),
]
for title, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + ": ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("Raster / Vector File Selection", level=4)
doc.add_paragraph(
    "Regardless of mode, a spatial data file must be specified for each layer. Two methods are available:"
)
items = [
    ("Native File Dialog / File Browser", "Click the \"\U0001f4c2 Choose File\" button. The system first attempts to open the OS native file picker (GET /api/native-file-dialog/). If the application is running headlessly (e.g., in Docker without a display), it falls back to the built-in FileBrowserModal. The browser shows the server directory structure, listing only .tif/.tiff and .shp files. Recent folders are stored in browser localStorage (up to 6 entries) for quick access."),
    ("Manual Path Entry", "A text field allows pasting the full file path manually (e.g., \"C:\\data\\solar_irradiation.tif\")."),
]
for title, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + ": ")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    "Shapefile (.shp) support: Shapefiles are accepted for Off-Shore Wind seabed substrate layers "
    "(EMODnet format). When a .shp file is selected, a note appears informing the user that it will "
    "be processed as a vector seabed layer using categorical substrate analysis."
)

add_note(doc, "All raster files must be in EPSG:3857 (Web Mercator) projection. The system validates the CRS on submission and rejects incompatible files. Maximum raster file size: 10 GB.")

doc.add_heading("Add Layer Button", level=4)
doc.add_paragraph(
    "Clicking the \"\u2795 Add Layer\" button validates the input and adds the layer to the configured "
    "layers list via POST /api/layers/add/. Duplicate layer names are rejected with an error."
)

doc.add_heading("2.2.3 Configured Layers List", level=3)

add_placeholder(doc, "[Screenshot: Configured layers list with checkboxes and batch controls]")

doc.add_paragraph(
    "Below the \"Add New Layer\" form, a \"Configured Layers\" section displays all currently added "
    "layers. The count is shown in the section title (e.g., \"Configured Layers (5)\")."
)

doc.add_paragraph("Batch selection controls appear above the list:")
items = [
    ("All", "Checks all layer checkboxes \u2014 all layers will be included in the next analysis run."),
    ("None", "Unchecks all layer checkboxes \u2014 no layers will be included until manually checked."),
]
for title, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + ": ")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph("Each layer card shows:")
items = [
    "A checkbox \u2014 checked layers are included in the analysis run; unchecked layers are skipped for the current run but remain configured",
    "An icon indicating predefined (\U0001f3f7\ufe0f tag) vs. custom (\U0001f527 gear) origin",
    "The layer name",
    "The assigned analysis modes (comma-separated)",
    "The filename of the associated spatial file",
    "A red \"\u00d7\" button to permanently remove the layer (calls DELETE /api/layers/<index>/remove/)",
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

add_tip(doc, "Use checkboxes to selectively re-run analysis on specific layers without re-processing all layers. Only checked layers are passed to Run Analysis.")

doc.add_heading("2.2.4 Layer Map Preview", level=3)
doc.add_paragraph(
    "An interactive Leaflet map (LayerMapPreview) below the configured layers list provides spatial context:"
)
items = [
    ("Boundary Layer", "The study area boundary polygon (toggleable)."),
    ("Grid Layer", "The grid cells overlay (toggleable)."),
    ("Raster Layer Previews", "Each configured raster layer can be toggled on/off individually. When toggled on for the first time, the system requests a server-rendered raster preview image (base64-encoded PNG with a blue-cyan-green-yellow-red colormap). The preview is overlaid on the map at 70% opacity."),
    ("Base Map Toggle", "Street (OpenStreetMap) or Satellite (ArcGIS) base maps."),
]
for title, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + ": ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("2.2.5 Run Analysis", level=3)
doc.add_paragraph(
    "The \"\U0001f680 Run Analysis\" button (green, full-width) starts the spatial analysis pipeline "
    "as an asynchronous task (POST /api/analysis/run-async/ with selected_indices). "
    "Only checked layers are included. The process:"
)
steps = [
    "A processing overlay appears with animated dots and a live progress message",
    "The backend loads the grid data and reprojects it to match each raster CRS",
    "For raster layers (.tif): the UniversalRasterScorer engine computes the requested metrics using memory-aware parallel chunking",
    "For vector seabed layers (.shp): the seabed_scorer performs spatial intersection to determine the dominant substrate category per cell",
    "Analysis modes are computed per cell: distance, coverage, mean, max, min, median, std, or categorical",
    "Results are merged into a single DataFrame and stored in the session as scoring_results.pkl",
    "The overlay disappears and results are displayed",
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

add_tip(doc, "Analysis time scales with cell count, layer count, and raster resolution. The system processes layers in parallel using multiple CPU workers. For large studies, use larger cell sizes to reduce processing time.")

doc.add_heading("2.2.6 Analysis Modes Reference", level=3)
add_table(doc,
    ["Mode", "Output Column Suffix", "Description"],
    [
        ["distance", "_dist_km", "Distance (km) from each cell to the nearest pixel matching the target value, using GDAL proximity analysis with large-raster downsampling"],
        ["coverage", "_coverage_pct", "Percentage (0\u2013100) of the cell's area covered by pixels matching the target value"],
        ["mean", "_mean", "Mean pixel value across all pixels within the cell boundary"],
        ["max", "_max", "Maximum pixel value within the cell"],
        ["min", "_min", "Minimum pixel value within the cell"],
        ["median", "_median", "Median pixel value within the cell"],
        ["std", "_std", "Standard deviation of pixel values within the cell"],
        ["categorical", "_dominant", "Dominant substrate category (Off-Shore seabed .shp layers only)"],
    ],
    col_widths=[2, 2.5, 10]
)

doc.add_heading("2.2.7 Analysis Results", level=3)

add_placeholder(doc, "[Screenshot: Analysis results with statistics cards and data table]")

doc.add_paragraph("After successful analysis, the results section displays:")

doc.add_heading("Success Banner", level=4)
doc.add_paragraph(
    "A green-bordered container with the success message and a \"\U0001f4e5 Download\" button "
    "for exporting the full results as raster_analysis_results.csv (semicolon-separated, "
    "comma decimals, UTF-8 BOM \u2014 compatible with Excel)."
)

doc.add_heading("Statistics Cards", level=4)
doc.add_paragraph(
    "A grid of summary statistic cards (up to 4 per row) \u2014 one for each computed column. "
    "Each card shows the column name, the average value (avg), and the min\u2013max range."
)

doc.add_heading("Data Table", level=4)
doc.add_paragraph(
    "A full-featured interactive AnalysisResultsTable (powered by TanStack React Table) displays the results:"
)
items = [
    ("Row Count Selector", "A dropdown to control rows per page (10, 25, 50, or 100)."),
    ("Sortable Columns", "Click any column header to sort ascending (\u25b2) or descending (\u25bc)."),
    ("Column Filters", "A text input below each column header enables real-time filtering."),
    ("Pagination", "Navigation buttons (first, previous, next, last) and a page indicator."),
    ("Number Formatting", "Numeric values are rounded to 3 decimal places."),
    ("Row Click", "Clicking any row focuses the corresponding grid cell on the map."),
]
for title, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + ": ")
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  2.3 TAB 3 — SCORING
# ════════════════════════════════════════════════════════════════════
doc.add_heading("2.3 Tab 3 \u2014 Scoring", level=2)

add_placeholder(doc, "[Screenshot: Scoring tab with layer configuration cards]")

doc.add_paragraph(
    "The Scoring tab is the third pipeline step. "
    "Raw metric values from Step 2 are converted into normalized scores (0\u2013100) "
    "using configurable scoring levels and weights. Hard exclusion constraints eliminate "
    "cells that violate critical thresholds. Transmission infrastructure layers "
    "(kV Lines and Substations) are handled via a dedicated \"Connection\" mode whose weights "
    "feed into cluster connection scoring in Step 4."
)
doc.add_paragraph(
    "When the tab is opened, it automatically loads the available columns from the "
    "Step 2 analysis results and groups them by layer name based on column suffixes "
    "(_dist_km, _coverage_pct, _mean, _max, _min, _median, _std, _dominant). "
    "A configuration card is rendered for each detected layer group."
)

doc.add_heading("2.3.1 Prerequisites", level=3)
doc.add_paragraph(
    "This tab requires analysis results from Step 2. If no analysis data is available, "
    "a warning message is displayed: \"No analysis data available. Complete Step 2 first.\""
)

doc.add_heading("2.3.2 Import CSV", level=3)
doc.add_paragraph(
    "A header button labeled \"Import CSV\" allows importing a previously saved scoring file. "
    "The uploaded CSV (via POST /api/scoring/import-csv/) must contain a cell_id column. "
    "This is useful to resume work from a previously exported scoring results file."
)

doc.add_heading("2.3.3 Weight Validation Banner", level=3)
doc.add_paragraph(
    "A persistent banner at the top of the configuration area displays the running total "
    "weight percentage across all active layers. The \"\U0001f680 Run Scoring\" button is disabled "
    "until the total equals exactly 100%. The banner provides:"
)
items = [
    "Current total weight % (green when 100%, red otherwise)",
    "\"Distribute Evenly\" button \u2014 divides 100% equally across all layers in Scoring or Connection mode, overwriting current weights",
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("2.3.4 Layer Configuration Cards", level=3)
doc.add_paragraph(
    "For each layer analyzed in Step 2, a configuration card is displayed. "
    "Four mode buttons in the card header control how the layer participates in scoring:"
)

add_table(doc,
    ["Mode", "Description"],
    [
        ["Scoring", "Layer metrics are scored using configurable level thresholds and contribute a weighted amount to FINAL_GRID_SCORE"],
        ["Connection", "Used for kV transmission layers (110/220/400 kV Lines and Substations). kV layers are auto-set to this mode. The configured weight is used in Step 4 Overall Score calculation, not in Step 3 FINAL_GRID_SCORE."],
        ["Exclusion", "Hard knockout constraint \u2014 cells exceeding the threshold receive FINAL_GRID_SCORE = 0 permanently"],
        ["Skip", "Layer is excluded from all scoring and connection calculations"],
    ],
    col_widths=[2.5, 11.5]
)

doc.add_heading("Scoring Mode \u2014 Layer Type: single_mode", level=4)
doc.add_paragraph(
    "Applies to layers with a single metric column (e.g., _mean, _max, or standalone _dist_km). "
    "Configuration options:"
)
items = [
    ("Weight (%)", "Percentage contribution (0\u2013100) of this layer to FINAL_GRID_SCORE."),
    ("Normalize by max", "Checkbox. When checked, all cell values are divided by the column maximum before applying level scoring, making scores relative to the study area's own range."),
    ("Scoring Levels (4 rows)", "Level 1 (best) through Level 4 (worst). Each level has Min, Max, and Score fields. Adjacent level boundaries are linked: changing one level's Min auto-updates the adjacent level's Max. A level is highlighted red with a \"min \u2265 max\" warning if invalid."),
]
for title, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + ": ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("Scoring Mode \u2014 Layer Type: distance_coverage", level=4)
doc.add_paragraph(
    "Applies to layers with both _dist_km and _coverage_pct columns (e.g., Agricultural Areas, Forest). "
    "Configuration options:"
)
items = [
    ("Weight (%)", "Percentage contribution to FINAL_GRID_SCORE."),
    ("Max Coverage Threshold (%)", "Default 5%. If a cell's coverage percentage exceeds this threshold, the distance score is set to the minimum level (worst) regardless of the actual distance value. This penalizes cells that already significantly overlap with the obstacle. The constraint preview displays the formula (e.g., \"Agricultural_Areas_coverage_pct \u2264 5%\")."),
    ("Distance Scoring Levels (4 rows)", "Four levels mapping distance ranges to scores."),
]
for title, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + ": ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("Scoring Mode \u2014 Layer Type: bathymetry_dual (Off-Shore Wind only)", level=4)
doc.add_paragraph(
    "Applies to depth-sensitive layers: Bathymetry, Wind Speed, and Slope in Off-Shore Wind mode. "
    "Accounts for two turbine foundation types depending on water depth:"
)
items = [
    ("Weight (%)", "Percentage contribution to FINAL_GRID_SCORE."),
    ("Depth Threshold (m)", "Default 60 m. Cells with depth \u2264 threshold use the Bottom Fixed level table; cells deeper than the threshold use the Floating level table (for floating-foundation turbines)."),
    ("Bottom Fixed Levels (4 rows)", "Score levels applied within the bottom-fixed depth range."),
    ("Floating Levels (4 rows)", "Score levels applied beyond the depth threshold."),
]
for title, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + ": ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("Scoring Mode \u2014 Layer Type: seabed_categorical (Off-Shore Wind only)", level=4)
doc.add_paragraph(
    "Applies to EMODnet seabed substrate layers processed from shapefile data. "
    "Cells are scored based on the dominant substrate category:"
)
add_table(doc,
    ["EMODnet Substrate", "Internal Category", "Default Score"],
    [
        ["Sand, seabed", "Sand", "100"],
        ["Coarse & mixed sediment", "Gravel", "70"],
        ["Fine mud, muddy sand, sandy mud", "Rack / Bad rack / Mud", "40"],
        ["Rock, boulder, Posidonia meadows", "Boulder / Stony / Silt", "0"],
    ],
    col_widths=[4, 3.5, 2.5]
)
add_note(doc, "Seabed scoring only applies to bottom-fixed cells (depth \u2264 depth threshold). Floating-depth cells receive score 0 from this layer type, as substrate is irrelevant for floating foundations.")

doc.add_heading("Connection Mode Configuration", level=4)
doc.add_paragraph(
    "kV layer cards (110/220/400 kV Lines and Substations) are automatically set to Connection mode. "
    "Only one configuration field is shown:"
)
p = doc.add_paragraph(style='List Bullet')
run = p.add_run("Weight (%): ")
run.bold = True
p.add_run(
    "The percentage weight assigned to this connection asset type. "
    "Used in Step 4 cluster scoring (Overall Score calculation), not in Step 3 FINAL_GRID_SCORE."
)

doc.add_heading("Exclusion Mode Configuration", level=4)
doc.add_paragraph(
    "When a layer is set to \"Exclusion\" mode:"
)
items = [
    ("Metric Selector", "If the layer has multiple analysis columns, a dropdown selects which column to use as the constraint metric."),
    ("Threshold Value", "A numeric input for the exclusion limit."),
    ("Mode Selector", "Specifies whether the constraint is a maximum (\u2264) or minimum (\u2265) threshold."),
    ("Constraint Preview", "Displays the constraint formula (e.g., \"Agricultural_Areas_coverage_pct \u2264 30\") for clarity."),
]
for title, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + ": ")
    run.bold = True
    p.add_run(desc)

p = doc.add_paragraph()
run = p.add_run("Warning: ")
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
p.add_run("Cells violating an exclusion constraint have FINAL_GRID_SCORE forced to 0. These cells are excluded from cluster analysis in Step 4.")

doc.add_heading("2.3.5 Run Scoring", level=3)
doc.add_paragraph(
    "The \"\U0001f680 Run Scoring\" button (purple, full-width) is enabled only when total weights sum to 100%. "
    "Clicking it starts scoring as an asynchronous task (POST /api/scoring/run-async/). The process:"
)
steps = [
    "Collects scoring configuration (levels, weights, layer types) and exclusion constraints from all layer cards",
    "For single_mode layers: maps metric values through 4 scoring levels to a 0\u2013100 score",
    "For distance_coverage layers: applies coverage threshold knockout, then distance levels",
    "For bathymetry_dual layers: routes cells to bottom-fixed or floating level tables based on depth",
    "For seabed_categorical layers: maps dominant substrate string to a configured score",
    "Applies layer weight to compute each layer's weighted contribution",
    "Sums all weighted contributions into FINAL_GRID_SCORE (0\u2013100)",
    "Applies hard exclusion constraints \u2014 violating cells set to FINAL_GRID_SCORE = 0",
    "Records exclusion reasons (EXCLUSION_REASONS column) and stores results as final_scored_results.pkl",
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_heading("2.3.6 Scoring Results", level=3)

add_placeholder(doc, "[Screenshot: Scoring results with distribution cards and preview table]")

doc.add_paragraph("After scoring completes:")

doc.add_heading("Score Distribution Cards", level=4)
add_table(doc,
    ["Category", "Score Range", "Color", "Description"],
    [
        ["Excellent", "\u2265 80", "Green", "Highly suitable cells \u2014 top candidates"],
        ["Good", "60 \u2013 79", "Light Green", "Well-suited cells with minor limitations"],
        ["Fair", "40 \u2013 59", "Yellow", "Moderate suitability"],
        ["Poor", "20 \u2013 39", "Orange", "Low suitability \u2014 significant constraints"],
        ["Very Poor", "1 \u2013 19", "Red", "Unsuitable \u2014 multiple unfavorable factors"],
        ["Excluded", "= 0", "Gray", "Hard-excluded cells via constraint thresholds"],
    ],
    col_widths=[2.5, 2, 2, 7.5]
)

doc.add_heading("Summary Statistics", level=4)
for item in ["Total: Total number of cells processed", "Excluded: Cells with FINAL_GRID_SCORE = 0 due to hard constraints", "Avg Score: Mean FINAL_GRID_SCORE across non-excluded cells"]:
    p = doc.add_paragraph(style='List Bullet')
    parts = item.split(": ", 1)
    run = p.add_run(parts[0] + ": ")
    run.bold = True
    p.add_run(parts[1])

doc.add_heading("Exclusion Summary Table", level=4)
doc.add_paragraph(
    "Lists each constraint layer with the column used, threshold applied, "
    "and number of cells excluded by that constraint."
)

doc.add_heading("Download & Data Preview", level=4)
doc.add_paragraph(
    "A \"\U0001f4e5 Download\" button exports the scored dataset via GET /api/scoring/download/ "
    "as final_scored_analysis.csv (semicolon-separated, comma decimals, UTF-8 BOM). "
    "An AnalysisResultsTable preview shows FINAL_GRID_SCORE, per-layer scores, and EXCLUSION_REASONS."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  2.4 TAB 4 — CLUSTER & AGGREGATION
# ════════════════════════════════════════════════════════════════════
doc.add_heading("2.4 Tab 4 \u2014 Cluster & Aggregation", level=2)

add_placeholder(doc, "[Screenshot: Cluster tab with four configuration sub-tabs]")

doc.add_paragraph(
    "The Cluster & Aggregation tab is the final pipeline step, combining spatial clustering, "
    "transmission connection scoring, and financial analysis. Adjacent high-scoring grid cells "
    "are grouped into project-sized clusters; each cluster is evaluated for its optimal grid "
    "connection and financial feasibility."
)
doc.add_paragraph(
    "The tab is organized into four configuration sub-tabs (Capacity Constraints, Connection Rules, "
    "Financial Constants, Technical Constants) followed by the Run button."
)

doc.add_heading("2.4.1 Prerequisites", level=3)
doc.add_paragraph(
    "This tab requires scored results from Step 3. If no scoring data is available, "
    "a warning is displayed: \"No scoring data available. Complete Step 3 first.\""
)
doc.add_paragraph(
    "Alternatively, a pre-scored CSV file can be uploaded directly via POST /api/cluster/upload-csv/ "
    "(must contain wkt and FINAL_GRID_SCORE columns; separator is auto-detected)."
)

doc.add_heading("2.4.2 Sub-Tab A: Capacity Constraints (\u26a1)", level=3)

add_placeholder(doc, "[Screenshot: Capacity Constraints sub-tab]")

add_table(doc,
    ["Parameter", "Input Type", "Range / Step", "Default", "Description"],
    [
        ["Nominal Capacity (MW)", "Number", "Min 0.1, Step 0.5", "13 MW", "Assumed power capacity per grid cell. For wind, this represents the turbine nameplate rating."],
        ["Max Cluster Capacity (MW)", "Number", "Min 10, Step 10", "250 MW", "Maximum total capacity per cluster. Over-capacity connected components are split via greedy BFS to maintain spatial contiguity."],
        ["Adjust capacity for coverage", "Checkbox", "On / Off", "Checked", "Reduces each cell's effective capacity proportionally by its coverage percentages. A cell with 30% obstacle coverage gets 70% of nominal capacity."],
    ],
    col_widths=[3.5, 2, 2.5, 1.5, 5.5]
)

doc.add_heading("2.4.3 Sub-Tab B: Connection Rules (\U0001f4cf)", level=3)

add_placeholder(doc, "[Screenshot: Connection Rules editable table]")

doc.add_paragraph(
    "An editable table of cluster connection scoring rules, pre-filtered to show only rules "
    "for kV layers configured in Step 2. If no kV layers were configured, a hint prompts "
    "the user to add kV layers in Step 2 to unlock connection scoring."
)

doc.add_paragraph("Each table row represents a scoring rule for a specific asset type and capacity range:")
add_table(doc,
    ["Column", "Description"],
    [
        ["Criteria", "Transmission asset name (read-only)"],
        ["Cap MW Range", "Cluster capacity range this rule applies to (0\u201330, 30\u201370, 70\u2013180, or 180\u2013400 MW)"],
        ["L1 Min / Max / Score", "Level 1 (closest/best): distance range (km) and assigned score"],
        ["L2 Min / Max / Score", "Level 2: distance range and score"],
        ["L3 Min / Max / Score", "Level 3: distance range and score"],
        ["L4 Min / Max / Score", "Level 4 (farthest/worst): distance range and score"],
    ],
    col_widths=[3.5, 10.5]
)
doc.add_paragraph(
    "The system evaluates each cluster's minimum distance to every configured transmission asset, "
    "applies the rules for the cluster's capacity range, and selects the best connection "
    "(highest score; ties broken by shortest distance). "
    "The \"\U0001f4be Save Rules\" button persists changes via PUT /api/scoring-rules/."
)
add_note(doc, "For Off-Shore Wind mode, only 220kV and 400kV assets are available. The rules table adjusts to show only offshore-relevant entries.")

doc.add_heading("2.4.4 Sub-Tab C: Financial Constants (\U0001f4b0)", level=3)

add_placeholder(doc, "[Screenshot: Financial Constants editing grid]")

doc.add_paragraph(
    "Editable financial parameters used for CAPEX and LCOE calculations. "
    "Fields are shown or hidden based on the active project mode."
)
add_table(doc,
    ["Parameter", "Default", "Applies To", "Description"],
    [
        ["PV CAPEX per MW", "$500,000", "Solar PV", "Capital cost per MW for PV installations"],
        ["Wind CAPEX per MW", "$1,000,000", "Wind", "Capital cost per MW for wind installations"],
        ["Substation Installation Cost Ratio", "Solar: 8%, Wind: 6%", "All", "Substation cost as fraction of generation CAPEX"],
        ["Line Expropriation Cost Ratio", "10%", "All", "Land expropriation cost for line corridor"],
        ["Land Expropriation Cost Ratio", "10%", "All", "Land acquisition as fraction of CAPEX"],
        ["Transport Network Base Cost", "$400,000", "Wind", "Fixed base cost for access infrastructure"],
        ["Transport Network Cost per MW", "$500 / MW", "Wind", "Variable transport cost per installed MW"],
    ],
    col_widths=[4, 2, 2, 6]
)

doc.add_paragraph("Transmission cost table (from financial_constants.json):")
add_table(doc,
    ["Type", "kV", "Capacity Range (MW)", "Cost/km", "Fixed Cost"],
    [
        ["Line", "110", "0 \u2013 30", "$170,000", "$0"],
        ["Line", "110", "30 \u2013 70", "$170,000", "$0"],
        ["Line", "220", "70 \u2013 180", "$280,000", "$0"],
        ["Line", "400", "180 \u2013 400", "$400,000", "$0"],
        ["Substation", "110", "0 \u2013 30", "$170,000", "$500,000"],
        ["Substation", "110", "30 \u2013 70", "$170,000", "$1,000,000"],
        ["Substation", "220", "70 \u2013 180", "$280,000", "$3,000,000"],
        ["Substation", "400", "180 \u2013 400", "$400,000", "$8,000,000"],
    ],
    col_widths=[2.5, 1.5, 3, 2.5, 2.5]
)
doc.add_paragraph(
    "The \"\U0001f4be Save Financial Constants\" button persists changes via PUT /api/financial-constants/."
)

doc.add_heading("2.4.5 Sub-Tab D: Technical Constants (\u2699\ufe0f)", level=3)

add_placeholder(doc, "[Screenshot: Technical Constants sub-tab]")

items = [
    ("Capacity Factor Override (0\u20131)", "Optional numeric input. If set, overrides the per-cluster calculated capacity factor for all clusters. If left empty, each cluster's CF is calculated from its energy yield and installed capacity."),
    ("Cp Values Table (Wind projects only)", "Editable scrollable table of wind speed (m/s) vs. power coefficient (Cp) pairs. Defines the turbine power curve. Matches the specific turbine model's aerodynamic performance. The \"\U0001f4be Save Technical Constants\" button persists changes via PUT /api/cp-values/."),
]
for title, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + ": ")
    run.bold = True
    p.add_run(desc)

add_tip(doc, "Populate the Cp table with the manufacturer's power curve for the turbine model under consideration. This significantly affects energy yield and LCOE accuracy.")

doc.add_heading("2.4.6 Run Cluster Analysis", level=3)
doc.add_paragraph(
    "The \"\U0001f680 Run Cluster Analysis\" button (indigo, full-width) initiates the complete pipeline "
    "as an asynchronous task (POST /api/cluster/run-async/). Three phases:"
)

doc.add_heading("Phase 1: Spatial Clustering (ClusterEngine)", level=4)
steps = [
    "Loads scored data, filters out cells with FINAL_GRID_SCORE \u2264 0",
    "Calculates effective capacity per cell (accounting for clipped edge cells and optional coverage reduction)",
    "Performs spatial adjacency analysis \u2014 identifies cells sharing a boundary via spatial intersection join",
    "Builds a NetworkX graph with cells as nodes and shared edges as graph edges",
    "Identifies connected components (candidate clusters)",
    "Enforces max cluster capacity: over-capacity components split via greedy BFS maintaining spatial contiguity",
    "Dissolves cell geometries into a unified cluster polygon per cluster",
    "Aggregates metrics: sum for capacity, mean for scores/coverages, min for all distance columns",
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_heading("Phase 2: Connection Scoring (ClusterScorer)", level=4)
steps = [
    "For each cluster, finds minimum distance to each configured transmission asset type across all member cells",
    "Evaluates each asset against connection rules for the cluster's capacity range",
    "Selects optimal connection: highest score; ties broken by shortest distance",
    "Records winning connection type, kV level, distance, and score",
    "Wind only: classifies cluster as \"Distribution\" (< 30 MW) or \"Transmission\" (\u2265 30 MW) grid",
    "Computes Overall Score = Mean Cell FINAL_GRID_SCORE + (connection weight \u00d7 connection score)",
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_heading("Phase 3: Financial Analysis (FinancialScorer)", level=4)

doc.add_heading("Solar PV Financial Model:", level=4)
for item in [
    "CAPEX OF PV = Capacity \u00d7 PV CAPEX per MW",
    "SUBSTATION COST = CAPEX OF PV \u00d7 Substation PV Ratio (8%)",
    "LAND COST = CAPEX OF PV \u00d7 Land Cost Ratio (10%)",
    "SLOPE COST = (CAPEX OF PV \u00d7 Mean Slope % \u00d7 9/15) / 100",
    "LINE CAPEX = Connection Distance \u00d7 Cost per km (+ Fixed Cost for substation connections)",
    "LINE EXPROPRIATION = LINE CAPEX \u00d7 Line Expropriation Ratio (10%)",
    "TOTAL CAPEX = PV + SUBSTATION + LAND + SLOPE + LINE + LINE_EXPROPRIATION",
    "Yearly Energy (MWh) = 1688 \u00d7 Solar Irradiation Fraction \u00d7 Capacity \u00b1 Temperature Correction",
    "Capacity Factor = Yearly Energy / (8,760 \u00d7 Capacity)",
    "LCOE ($/MWh) = (Annualized CAPEX + Annual OPEX) / Yearly Energy  [25-yr lifetime, 8% discount, 2% OPEX]",
    "Payback Period (Yrs) = TOTAL CAPEX / (Yearly Energy \u00d7 $50/MWh)",
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("Wind Financial Model (On-Shore & Off-Shore):", level=4)
for item in [
    "CAPEX = Capacity \u00d7 Wind CAPEX per MW",
    "SUBSTATION COST = CAPEX \u00d7 Substation Wind Ratio (6%)",
    "LAND COST = CAPEX \u00d7 Land Cost Ratio (10%)",
    "TRANSPORT NETWORKS = (Transport Distance \u00d7 $400,000) + (MW / 4 \u00d7 $2,000)",
    "LINE CAPEX = Connection Distance \u00d7 Cost per km (+ Fixed Cost for substation connections)",
    "LINE EXPROPRIATION = LINE CAPEX \u00d7 Line Expropriation Ratio (10%)",
    "TOTAL CAPEX = CAPEX + SUBSTATION + LAND + TRANSPORT + LINE + LINE_EXPROPRIATION",
    "Air Density (\u03c1) = 1.225 \u2212 0.264 \u00d7 (Altitude / 2000)  [elevation-based density correction]",
    "Cp = nearest-neighbor lookup from Cp values table (wind speed \u2192 power coefficient)",
    "Swept Area = \u03c0 \u00d7 rotor_radius\u00b2",
    "Yearly Energy (MWh) = (0.5 \u00d7 \u03c1 \u00d7 Area \u00d7 v\u00b3 \u00d7 Cp \u00d7 8,760 / 1,000,000) \u00d7 (MW / 4.2)",
    "LCOE ($/MWh) and Payback Period calculated using same assumptions as Solar",
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("Additional Output Metrics:", level=4)
for item in [
    "CAPEX/MW ($) = TOTAL CAPEX / Installed Capacity",
    "Scaled Overall Score = Overall Score \u00d7 (Capacity Factor / Max CF across all clusters)",
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("2.4.7 Cluster Results", level=3)

add_placeholder(doc, "[Screenshot: Cluster results with statistics, map, and data table]")

doc.add_heading("Success Banner & Download", level=4)
doc.add_paragraph(
    "A green banner with a \"\U0001f4e5 Download\" button exports clustered_scored_results.csv "
    "(semicolon-separated, comma decimals, UTF-8 BOM; geometry column excluded, wkt preserved)."
)

doc.add_heading("Summary Statistics Cards", level=4)
add_table(doc,
    ["Metric", "Description"],
    [
        ["Total Clusters", "Total number of clusters formed from eligible cells"],
        ["Avg Capacity (MW)", "Average installed capacity per cluster"],
        ["Total Capacity (MW)", "Sum of all cluster capacities"],
        ["Avg Overall Score", "Mean Overall Score across all clusters"],
        ["Avg LCOE ($/MWh)", "Average Levelized Cost of Energy"],
    ],
    col_widths=[3.5, 10.5]
)

doc.add_heading("Connection Distribution Cards", level=4)
items = [
    ("Connection Type Distribution", "Count and percentage of clusters grouped by winning connection type (Line vs. Substation)."),
    ("kV Distribution", "Count and percentage of clusters grouped by voltage level (110kV, 220kV, 400kV)."),
]
for title, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title + ": ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("Cluster Map (ClusterMapPreview)", level=4)
add_placeholder(doc, "[Screenshot: Cluster map with colored polygon overlays]")
doc.add_paragraph(
    "An interactive Leaflet map displays all cluster polygons as colored overlays. "
    "Each polygon represents the dissolved geometry of all member cells. "
    "Clicking a cluster polygon highlights it and scrolls the data table to that row. "
    "Layer toggles and base map selection are available."
)

doc.add_heading("Cluster Data Table", level=4)
doc.add_paragraph(
    "An AnalysisResultsTable shows per-cluster details: cluster ID, installed capacity (MW), "
    "number of member cells (Within_Cells_Count), best connection type/kV/distance/score, "
    "overall score, total CAPEX, CAPEX/MW, yearly energy (MWh), LCOE ($/MWh), "
    "capacity factor, payback period (years), and component cost breakdowns. "
    "Supports sorting, filtering, pagination, and row-click map focus."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  SECTION 3 — PROJECT MODE DIFFERENCES
# ════════════════════════════════════════════════════════════════════
doc.add_heading("3. Project Mode Differences", level=1)

doc.add_paragraph(
    "While the four-tab pipeline structure is identical across all project modes, "
    "each mode has distinct parameters, layers, and scoring configurations. "
    "This section summarizes the key differences."
)

doc.add_heading("3.1 Grid Parameters", level=2)
add_table(doc,
    ["Parameter", "Solar PV", "On-Shore Wind", "Off-Shore Wind"],
    [
        ["Grid Input", "Width \u00d7 Height (m)", "Turbine Diameter \u2192 3D \u00d7 5D", "Turbine Diameter \u2192 3D \u00d7 5D"],
        ["Turbine Diameter Range", "N/A", "20 \u2013 200 m, step 10", "20 \u2013 500 m, step 10"],
        ["Default Cell Size", "1,000 m \u00d7 1,000 m", "600 m \u00d7 1,000 m (200 m turbine)", "600 m \u00d7 1,000 m (200 m turbine)"],
        ["Boundary Source", "Country (NUTS) or Albania ADM", "Country (NUTS) or Albania ADM", "EEZ (Maritime Zones)"],
    ],
    col_widths=[3.5, 3.5, 3.5, 4]
)

doc.add_heading("3.2 Predefined Layers", level=2)
add_table(doc,
    ["Layer Category", "Solar PV", "On-Shore Wind", "Off-Shore Wind"],
    [
        ["Wind Resources", "\u2014", "Wind Speed (max, mean, min)", "Wind Speed (max, mean, min) \u2014 bathymetry_dual"],
        ["Solar Resources", "Solar Irradiation (mean), Temperature (mean)", "\u2014", "\u2014"],
        ["Terrain", "Slope (%) min/max/mean", "Slope (%), Altitude (max)", "Slope (%) bathymetry_dual; Bathymetry (max) bathymetry_dual"],
        ["Seabed", "\u2014", "\u2014", "Sea bed (.shp EMODnet) seabed_categorical"],
        ["Grid Infrastructure", "110/220/400 kV Lines & Substations (distance)", "110/220/400 kV Lines & Substations (distance)", "220/400 kV Lines & Substations only (distance)"],
        ["Land/Marine Use", "Agriculture, Forest, Urban, Military, Protected (dist+cov)", "Agriculture, Airports, Forest, Land Use, Military, Protected (dist+cov)", "Fishing, Military, Shipping, Tourism, Protected (dist+cov)"],
        ["Transportation", "Transport Networks (distance)", "Transport Networks (distance)", "Ports (distance), Subsea Cables (distance)"],
        ["Natural Resources", "Energy Sources, Hydrography, Minerals", "Energy, Hydrography, Minerals, Natural Risk", "Natural Risk"],
    ],
    col_widths=[3, 3.5, 3.5, 4]
)

doc.add_heading("3.3 Scoring Types Used per Mode", level=2)
add_table(doc,
    ["Scoring Type", "Solar PV", "On-Shore Wind", "Off-Shore Wind"],
    [
        ["single_mode", "Yes \u2014 Slope, Irradiation, Temperature", "Yes \u2014 Wind, Altitude, Slope", "Yes \u2014 some marine layers"],
        ["distance_coverage", "Yes \u2014 Agriculture, Forest, Urban, etc.", "Yes \u2014 Agriculture, Airport, Forest, etc.", "Yes \u2014 Fishing, Military, Shipping, etc."],
        ["bathymetry_dual", "No", "No", "Yes \u2014 Bathymetry, Wind Speed, Slope"],
        ["seabed_categorical", "No", "No", "Yes \u2014 Sea bed substrate"],
    ],
    col_widths=[3.5, 3.5, 3.5, 4]
)

doc.add_heading("3.4 Color Theme", level=2)
add_table(doc,
    ["Element", "Solar PV", "On-Shore Wind", "Off-Shore Wind"],
    [
        ["Header Accent", "Orange", "Dark Blue (Navy)", "Cyan / Blue"],
        ["Card Border", "Orange-300", "Blue-400", "Cyan-300"],
        ["Selection Button", "Orange", "Dark Blue", "Blue"],
    ],
    col_widths=[3.5, 4, 4, 4]
)

doc.add_heading("3.5 Financial Model Differences", level=2)
add_table(doc,
    ["Parameter", "Solar PV", "Wind (On-Shore & Off-Shore)"],
    [
        ["Base CAPEX", "$500,000 / MW", "$1,000,000 / MW"],
        ["Substation Ratio", "8% of CAPEX", "6% of CAPEX"],
        ["Slope Cost", "Included (terrain correction)", "Not applicable"],
        ["Transport Cost", "Not applicable", "Included (access road infrastructure)"],
        ["Energy Calculation", "1688 \u00d7 irradiance fraction \u00d7 capacity \u00b1 temperature correction", "0.5 \u00d7 \u03c1 \u00d7 swept area \u00d7 v\u00b3 \u00d7 Cp \u00d7 8,760 \u00d7 (MW/4.2)"],
        ["Air Density Correction", "No", "Yes \u2014 \u03c1 = 1.225 \u2212 0.264 \u00d7 (altitude/2000)"],
        ["Capacity Factor Source", "Solar irradiation-based", "Wind speed & Cp table-based"],
    ],
    col_widths=[3, 5.5, 5.5]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  SECTION 4 — APPENDICES
# ════════════════════════════════════════════════════════════════════
doc.add_heading("4. Appendices", level=1)

doc.add_heading("4.1 API Endpoints Reference", level=2)
doc.add_paragraph(
    "All API endpoints are served at http://localhost:8000/api/. "
    "All requests requiring session context must include the X-Session-ID header "
    "(UUID stored in browser localStorage as dashboard_session_id)."
)

add_table(doc,
    ["Category", "Method", "Endpoint", "Description"],
    [
        ["Project", "POST", "/project/select/", "Create session: project_type = Solar | OnShore | OffShore"],
        ["Project", "GET", "/project/status/", "Get session state and step completion flags"],
        ["Project", "POST", "/project/reset/", "Clear session; keep_project_type: bool"],
        ["Project", "GET", "/project/config/", "Get layer/scoring config for active project type"],
        ["Grid", "POST", "/grid/create/", "Create grid from country/EEZ/Albania boundary"],
        ["Grid", "POST", "/grid/upload/", "Upload existing grid CSV"],
        ["Grid", "GET", "/grid/download/", "Download grid.csv"],
        ["Grid", "GET", "/countries/", "List NUTS European countries"],
        ["Grid", "GET", "/eez-zones/", "List EEZ maritime zones"],
        ["Grid", "GET", "/albania/regions/", "List Albania ADM1 regions"],
        ["Grid", "GET", "/albania/districts/", "List Albania ADM2 districts (filter: ?region=...)"],
        ["Layers", "POST", "/layers/add/", "Add raster (.tif) or vector (.shp) layer"],
        ["Layers", "DELETE", "/layers/<index>/remove/", "Remove layer by index"],
        ["Analysis", "POST", "/analysis/run-async/", "Start async raster/vector analysis (returns task_id)"],
        ["Analysis", "GET", "/analysis/download/", "Download raster_analysis_results.csv"],
        ["Scoring", "POST", "/scoring/run-async/", "Start async level scoring (returns task_id)"],
        ["Scoring", "GET", "/scoring/download/", "Download final_scored_analysis.csv"],
        ["Scoring", "POST", "/scoring/import-csv/", "Import pre-computed scoring CSV"],
        ["Cluster", "POST", "/cluster/run-async/", "Start async cluster analysis (returns task_id)"],
        ["Cluster", "GET", "/cluster/download/", "Download clustered_scored_results.csv"],
        ["Cluster", "POST", "/cluster/upload-csv/", "Upload pre-scored CSV as cluster input"],
        ["Config", "GET/PUT", "/financial-constants/", "Read / write financial parameters"],
        ["Config", "GET/PUT", "/cp-values/", "Read / write wind turbine Cp table"],
        ["Config", "GET/PUT", "/scoring-rules/", "Read / write cluster connection scoring rules"],
        ["Tasks", "GET", "/task/<task_id>/progress/", "Poll async task progress: {progress, message, status}"],
        ["Browse", "GET", "/browse/", "Browse server filesystem for .tif / .shp files"],
        ["Browse", "GET", "/native-file-dialog/", "Open OS native file picker dialog"],
    ],
    col_widths=[2, 1.5, 4.5, 6]
)

doc.add_heading("4.2 Glossary of Terms", level=2)
add_table(doc,
    ["Term", "Definition"],
    [
        ["CAPEX", "Capital Expenditure \u2014 total upfront investment cost"],
        ["Cp (Power Coefficient)", "Fraction of wind energy extracted by a turbine (0 to Betz limit ~0.59). Defined via lookup table."],
        ["CRS / EPSG", "Coordinate Reference System / European Petroleum Survey Group code. EPSG:3857 (Web Mercator) required for all input rasters."],
        ["EEZ", "Exclusive Economic Zone \u2014 maritime area (200 NM) with coastal state resource rights. Used as boundary for Off-Shore Wind mode."],
        ["EMODnet", "European Marine Observation and Data Network \u2014 provider of seabed substrate shapefile data."],
        ["GeoJSON", "Open standard format for encoding geographic data structures (geometries + attributes)."],
        ["Grid Cell", "Single rectangular polygon unit identified by a unique cell_id."],
        ["LCOE", "Levelized Cost of Energy \u2014 total lifetime cost per MWh produced."],
        ["NUTS", "Nomenclature of Territorial Units for Statistics \u2014 EU administrative division standard."],
        ["Raster", "Pixel-based spatial data format (GeoTIFF) for continuous variables."],
        ["Session ID", "UUID per project session, stored in localStorage and sent in X-Session-ID header."],
        ["Task ID", "UUID returned by async endpoints; used to poll /api/task/<id>/progress/."],
        ["WKT", "Well-Known Text \u2014 text representation of vector geometries."],
        ["Shapefile", "Geospatial vector format (.shp + companion files). Supported for seabed substrate layers."],
    ],
    col_widths=[3, 11]
)

doc.add_heading("4.3 Data Requirements", level=2)
doc.add_heading("Input Raster Files (.tif / .tiff)", level=3)
for item in [
    "Format: GeoTIFF (.tif / .tiff)",
    "Coordinate Reference System: EPSG:3857 (Web Mercator) \u2014 mandatory",
    "Maximum file size: 10 GB",
    "Recommended resolution: match or finer than the intended grid cell size",
    "Storage: place files in the data/ directory or specify the full server-side path during layer configuration",
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("Input Vector Files (.shp \u2014 Off-Shore Seabed only)", level=3)
for item in [
    "Format: ESRI Shapefile (.shp with companion .dbf, .shx, .prj files)",
    "Data source: EMODnet seabed substrate classification",
    "Must contain substrate classification attributes compatible with EMODnet schema",
    "The seabed scorer performs spatial intersection with grid cells and assigns dominant substrate category",
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("Upload Grid CSV Columns", level=3)
add_table(doc,
    ["Column", "Type", "Required", "Description"],
    [
        ["cell_id", "Integer", "Yes", "Unique identifier for each cell"],
        ["wkt", "String (WKT)", "Yes", "Cell geometry in Well-Known Text format (EPSG:3857)"],
    ],
    col_widths=[2.5, 3, 2, 6.5]
)

doc.add_heading("Download Output File Formats", level=3)
add_table(doc,
    ["Download", "Filename", "Separator", "Decimal", "Encoding"],
    [
        ["Grid", "grid.csv", "Comma (,)", "Dot (.)", "UTF-8"],
        ["Raster Analysis", "raster_analysis_results.csv", "Semicolon (;)", "Comma (,)", "UTF-8 BOM"],
        ["Scored Results", "final_scored_analysis.csv", "Semicolon (;)", "Comma (,)", "UTF-8 BOM"],
        ["Cluster Results", "clustered_scored_results.csv", "Semicolon (;)", "Comma (,)", "UTF-8 BOM"],
    ],
    col_widths=[2.5, 4.5, 2.5, 2, 2.5]
)
add_note(doc, "UTF-8 BOM encoding ensures correct display of special characters when opening CSV files in Microsoft Excel on Windows.")

doc.add_heading("4.4 Tips and Best Practices", level=2)
for item in [
    "Use the file browser's \"Recent\" button to quickly navigate to the last 6 visited directories",
    "The file browser persists the last directory across sessions via POST /api/browse/save-last-dir/",
    "Column headers in all data tables are clickable for sorting \u2014 click once ascending, again descending",
    "Use column filter inputs to search for specific cells or clusters by any attribute value",
    "The map supports standard Leaflet controls: scroll wheel to zoom, click-drag to pan",
    "Switch Street/Satellite base maps with the radio buttons in the map's top-right corner",
    "Download results at each step as a checkpoint \u2014 protects work against Docker restart or session loss",
    "Use \"Distribute Evenly\" in the Scoring tab to quickly set equal weights, then fine-tune individually",
    "Use layer checkboxes in Step 2 to selectively re-run analysis on specific layers only",
    "Scoring level boundaries are linked \u2014 changing one level's Min auto-updates the adjacent Max",
    "Review the Score Distribution cards after Step 3 \u2014 if too many cells are Excluded, loosen exclusion thresholds",
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("4.5 Troubleshooting", level=2)
add_table(doc,
    ["Issue", "Possible Cause", "Solution"],
    [
        ["\"No grid data available\" warning", "Step 1 not completed", "Go to Gridization tab and create or upload a grid"],
        ["Raster file rejected with CRS error", "File is not in EPSG:3857", "Reproject using QGIS or gdalwarp before adding it"],
        ["Run Analysis button disabled", "No layers checked for inclusion", "Use the All link or check individual layer checkboxes"],
        ["Run Scoring button disabled", "Total weights \u2260 100%", "Use \"Distribute Evenly\" or adjust weights manually until total = 100%"],
        ["Analysis runs very slowly", "Large study area or high-resolution rasters", "Use larger grid cells or lower-resolution rasters to reduce processing time"],
        ["Session data lost after restart", "Docker container was restarted", "Download CSV at each step to preserve work; re-upload in Step 4 if needed"],
        ["Map preview not loading", "Browser rendering issue", "Refresh the page (Ctrl+Shift+R) or clear browser cache"],
        ["File browser shows empty directory", "Wrong path or Docker volume not mounted", "Ensure data/ directory is mounted in Docker Compose and contains .tif / .shp files"],
        ["Albania sub-district dropdown not visible", "Region not selected yet", "Select a Region first; the District dropdown appears after region selection"],
        ["Shapefile returns empty seabed results", "Missing companion files or CRS mismatch", "Ensure .shp, .dbf, .shx files are all present; reproject if needed"],
        ["Cluster results: 0 clusters formed", "All cells excluded (score = 0)", "Loosen exclusion constraints in Step 3; thresholds may be too strict"],
        ["Connection Rules table is empty", "No kV layers in Step 2", "Add at least one kV transmission layer in Step 2 to enable connection scoring"],
        ["Scoring weights won't reach 100%", "Rounding in manual entry", "Use \"Distribute Evenly\", then fine-tune individual layer weights"],
    ],
    col_widths=[3.5, 3.5, 7]
)

# ── Save ─────────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "User_Manual.docx")
doc.save(output_path)
print(f"User manual generated successfully: {output_path}")
